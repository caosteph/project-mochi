"""Document generation — PDFs (reportlab) and Word docs (python-docx), both pure-python
(no native deps). Input is light Markdown-ish text: `#`/`##`/`###` headings, `-`/`*`
bullets, blank-line-separated paragraphs.
"""

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


def _blocks(content: str):
    """Yield (kind, text) for each line: heading level, bullet, blank, or paragraph."""
    for raw in (content or "").splitlines():
        line = raw.rstrip()
        stripped = line.lstrip()
        if not stripped:
            yield ("blank", "")
        elif stripped.startswith("### "):
            yield ("h3", stripped[4:])
        elif stripped.startswith("## "):
            yield ("h2", stripped[3:])
        elif stripped.startswith("# "):
            yield ("h1", stripped[2:])
        elif stripped[:2] in ("- ", "* "):
            yield ("bullet", stripped[2:])
        else:
            yield ("p", line)


def generate_pdf(title: str, content: str, path: Path) -> Path:
    styles = getSampleStyleSheet()
    hmap = {"h1": "Heading1", "h2": "Heading2", "h3": "Heading3"}
    flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    bullets: list = []

    def flush_bullets():
        if bullets:
            flow.append(ListFlowable([ListItem(Paragraph(b, styles["BodyText"])) for b in bullets], bulletType="bullet"))
            bullets.clear()

    for kind, text_ in _blocks(content):
        if kind == "bullet":
            bullets.append(text_)
            continue
        flush_bullets()
        if kind in hmap:
            flow.append(Paragraph(text_, styles[hmap[kind]]))
        elif kind == "p":
            flow.append(Paragraph(text_, styles["BodyText"]))
        elif kind == "blank":
            flow.append(Spacer(1, 8))
    flush_bullets()
    SimpleDocTemplate(str(path), pagesize=letter, title=title).build(flow)
    return path


def generate_docx(title: str, content: str, path: Path) -> Path:
    doc = Document()
    doc.add_heading(title, level=0)
    levels = {"h1": 1, "h2": 2, "h3": 3}
    for kind, text_ in _blocks(content):
        if kind in levels:
            doc.add_heading(text_, level=levels[kind])
        elif kind == "bullet":
            doc.add_paragraph(text_, style="List Bullet")
        elif kind == "p":
            doc.add_paragraph(text_)
    doc.save(str(path))
    return path
